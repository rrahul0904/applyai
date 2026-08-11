from __future__ import annotations

import io
import re
import uuid
from datetime import date, datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches, Pt
from fastapi import APIRouter, Depends, HTTPException, Response
from pydantic import BaseModel
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.api.application_agent import _execution_payload, _owned_execution
from app.application_agent_models import ApplicationExecution
from app.career_models import CoverLetter, ResumeTailoring, ResumeTailoringRevision
from app.core.auth import get_current_user
from app.core.database import get_session
from app.core.internal_auth import require_internal_api
from app.core.storage import ObjectStorageProvider, get_object_storage
from app.models import CandidateEducation, CandidateExperience, CandidateProfile, CandidateSkill, JobSkill, User


router = APIRouter(prefix="/application-agent", tags=["application agent documents"])
internal_router = APIRouter(
    prefix="/internal/application-agent",
    tags=["internal-application-agent-documents"],
    dependencies=[Depends(require_internal_api)],
)

DOCX_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class DocumentReviewWrite(BaseModel):
    candidate_verified: bool = True


def _now() -> datetime:
    return datetime.now(timezone.utc)


def _date_label(value: date | None) -> str:
    return value.strftime("%b %Y") if value else ""


def _clean_filename(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "-", value.strip()).strip("-.")
    return cleaned or "application-document"


def _configure_doc(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    for style_name in ("Title", "Heading 1", "Heading 2"):
        style = document.styles[style_name]
        style.font.name = "Arial"


def _add_bullets(document: Document, text: str | None) -> None:
    if not text:
        return
    chunks = [item.strip(" •\t-") for item in re.split(r"\n+|(?<=\.)\s+(?=[A-Z])", text) if item.strip()]
    for chunk in chunks:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(chunk)


def _resume_bytes(session: Session, execution: ApplicationExecution, user: User) -> bytes:
    profile = session.scalar(select(CandidateProfile).where(CandidateProfile.user_id == user.id))
    if profile is None:
        raise HTTPException(status_code=409, detail="Candidate profile is required before generating a resume")

    experiences = list(
        session.scalars(
            select(CandidateExperience)
            .where(CandidateExperience.profile_id == profile.id)
            .order_by(CandidateExperience.start_date.desc().nullslast())
        )
    )
    education = list(
        session.scalars(
            select(CandidateEducation)
            .where(CandidateEducation.profile_id == profile.id)
            .order_by(CandidateEducation.end_date.desc().nullslast())
        )
    )
    skills = list(session.scalars(select(CandidateSkill).where(CandidateSkill.profile_id == profile.id)))
    job_skills = {
        value.lower()
        for value in session.scalars(select(JobSkill.normalized_name).where(JobSkill.job_id == execution.job_id))
    }
    skills.sort(key=lambda skill: (skill.normalized_name not in job_skills, skill.name.lower()))

    approved_highlights: list[str] = []
    resume_meta = dict((execution.documents or {}).get("resume") or {})
    artifact_id = resume_meta.get("artifact_id")
    if artifact_id:
        tailoring = session.scalar(
            select(ResumeTailoring).where(ResumeTailoring.artifact_id == uuid.UUID(str(artifact_id)))
        )
        if tailoring is not None:
            revisions = list(
                session.scalars(
                    select(ResumeTailoringRevision)
                    .where(
                        ResumeTailoringRevision.tailoring_id == tailoring.id,
                        ResumeTailoringRevision.candidate_decision.in_(["APPROVED", "EDITED"]),
                    )
                    .order_by(ResumeTailoringRevision.position)
                )
            )
            approved_highlights = [
                (revision.candidate_text or revision.suggested_text).strip()
                for revision in revisions
                if (revision.candidate_text or revision.suggested_text).strip()
            ]

    document = Document()
    _configure_doc(document)
    title = document.add_paragraph()
    title.alignment = 1
    run = title.add_run(" ".join(part for part in (user.first_name, user.last_name) if part) or user.email)
    run.bold = True
    run.font.size = Pt(18)
    contact = document.add_paragraph()
    contact.alignment = 1
    contact.add_run(user.email)

    if profile.current_title or profile.headline:
        heading = document.add_paragraph()
        heading.alignment = 1
        heading.add_run(profile.current_title or profile.headline or "").bold = True

    if profile.summary:
        document.add_heading("Professional Summary", level=1)
        document.add_paragraph(profile.summary)

    if skills:
        document.add_heading("Skills", level=1)
        document.add_paragraph(" • ".join(skill.name for skill in skills))

    if approved_highlights:
        document.add_heading("Relevant Highlights", level=1)
        for highlight in approved_highlights:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(highlight)

    if experiences:
        document.add_heading("Experience", level=1)
        for experience in experiences:
            row = document.add_paragraph()
            lead = row.add_run(f"{experience.title} — {experience.company_name}")
            lead.bold = True
            dates = " – ".join(filter(None, (_date_label(experience.start_date), _date_label(experience.end_date) or "Present")))
            if dates:
                row.add_run(f"  |  {dates}")
            _add_bullets(document, experience.description)

    if education:
        document.add_heading("Education", level=1)
        for item in education:
            row = document.add_paragraph()
            degree = ", ".join(part for part in (item.degree, item.field_of_study) if part)
            row.add_run(degree or "Education").bold = True
            row.add_run(f" — {item.institution}")
            if item.end_date:
                row.add_run(f"  |  {_date_label(item.end_date)}")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _cover_letter_bytes(execution: ApplicationExecution) -> bytes:
    cover = dict((execution.documents or {}).get("cover_letter") or {})
    body = str(cover.get("body") or "").strip()
    if not body:
        raise HTTPException(status_code=409, detail="Cover letter is not available")
    document = Document()
    _configure_doc(document)
    for block in re.split(r"\n\s*\n", body):
        if block.strip():
            document.add_paragraph(block.strip())
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _put_document(
    *,
    storage: ObjectStorageProvider,
    key: str,
    content: bytes,
) -> dict[str, Any]:
    storage.put(key=key, content=io.BytesIO(content), content_type=DOCX_TYPE)
    return {
        "storage_key": key,
        "content_type": DOCX_TYPE,
        "size": len(content),
        "generated_at": _now().isoformat(),
    }


@router.post("/executions/{execution_id}/documents/generate")
def generate_documents(
    execution_id: uuid.UUID,
    user: User = Depends(get_current_user),
    session: Session = Depends(get_session),
    storage: ObjectStorageProvider = Depends(get_object_storage),
) -> dict[str, Any]:
    execution = _owned_execution(session, user, execution_id)
    if execution.state in {"BROWSER_RUNNING", "CONFIRMED"}:
        raise HTTPException(status_code=409, detail="Application documents can no longer be regenerated")

    current = {key: dict(value) for key, value in (execution.documents or {}).items()}
    resume_bytes = _resume_bytes(session, execution, user)
    resume_filename = _clean_filename(f"{user.first_name or 'candidate'}-{user.last_name or ''}-tailored-resume.docx")
    resume_key = f"candidate/{user.id}/applications/{execution.application_id}/{execution.id}/resume.docx"
    resume = {
        **dict(current.get("resume") or {}),
        **_put_document(storage=storage, key=resume_key, content=resume_bytes),
        "filename": resume_filename,
        "candidate_verified": True,
        "truth_policy": "CANONICAL_PROFILE_PLUS_CANDIDATE_APPROVED_REVISIONS",
    }

    cover = dict(current.get("cover_letter") or {})
    if cover.get("body"):
        cover_bytes = _cover_letter_bytes(execution)
        cover_key = f"candidate/{user.id}/applications/{execution.application_id}/{execution.id}/cover-letter.docx"
        cover = {
            **cover,
            **_put_document(storage=storage, key=cover_key, content=cover_bytes),
            "filename": "cover-letter.docx",
        }
    execution.documents = {**current, "resume": resume, "cover_letter": cover}
    session.commit()
    session.refresh(execution)
    return _execution_payload(execution)


@router.patch("/executions/{execution_id}/documents/{document_type}")
def review_document(
    execution_id: uuid.UUID,
    document_type: str,
    body: DocumentReviewWrite,
    user: User = Depends(get_current_user),
    session: Session = Depends(get_session),
) -> dict[str, Any]:
    if document_type not in {"resume", "cover_letter"}:
        raise HTTPException(status_code=404, detail="Application document not found")
    execution = _owned_execution(session, user, execution_id)
    documents = {key: dict(value) for key, value in (execution.documents or {}).items()}
    document = dict(documents.get(document_type) or {})
    if not document:
        raise HTTPException(status_code=404, detail="Application document not found")
    document["candidate_verified"] = body.candidate_verified
    document["reviewed_at"] = _now().isoformat()
    documents[document_type] = document
    execution.documents = documents

    if document_type == "cover_letter" and document.get("id"):
        cover = session.get(CoverLetter, uuid.UUID(str(document["id"])))
        if cover is not None and cover.user_id == user.id:
            cover.candidate_verified = body.candidate_verified
    session.commit()
    session.refresh(execution)
    return _execution_payload(execution)


@internal_router.get("/executions/{execution_id}/documents/{document_type}")
def download_document_for_worker(
    execution_id: uuid.UUID,
    document_type: str,
    session: Session = Depends(get_session),
    storage: ObjectStorageProvider = Depends(get_object_storage),
) -> Response:
    if document_type not in {"resume", "cover_letter"}:
        raise HTTPException(status_code=404, detail="Application document not found")
    execution = session.get(ApplicationExecution, execution_id)
    if execution is None:
        raise HTTPException(status_code=404, detail="Application execution not found")
    document = dict((execution.documents or {}).get(document_type) or {})
    storage_key = document.get("storage_key")
    if not storage_key:
        raise HTTPException(status_code=409, detail="Application document has not been generated")
    if not document.get("candidate_verified"):
        raise HTTPException(status_code=409, detail="Candidate verification is required before document upload")
    try:
        content = storage.get(key=str(storage_key))
    except Exception as exc:
        raise HTTPException(status_code=404, detail="Application document object not found") from exc
    filename = Path(str(document.get("filename") or f"{document_type}.docx")).name
    return Response(
        content=content,
        media_type=str(document.get("content_type") or DOCX_TYPE),
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
