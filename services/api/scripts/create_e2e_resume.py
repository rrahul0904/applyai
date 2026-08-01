import argparse
from pathlib import Path

from docx import Document


def build_resume(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_paragraph("E2E Candidate")
    document.add_paragraph("Senior Data Engineer")
    document.add_paragraph("e2e.candidate@example.test")
    document.add_paragraph("Experience")
    document.add_paragraph("Senior Data Engineer | Example Labs | 2021 - Present")
    document.add_paragraph(
        "Built reliable production data platforms using Python, SQL, PostgreSQL, AWS, "
        "workflow orchestration, observability, automated testing, deployment pipelines, "
        "security controls, data governance, incident response, and cost optimization."
    )
    document.add_paragraph("Education")
    document.add_paragraph("Bachelor of Science, Computer Science | Example University | 2020")
    document.add_paragraph("Skills")
    document.add_paragraph("Python, SQL, AWS, PostgreSQL, Airflow, Docker, Data Engineering")
    document.add_paragraph("Certifications")
    document.add_paragraph("AWS Certified Data Engineer")
    document.save(path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Create the deterministic ApplyAI E2E resume.")
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    build_resume(args.output)


if __name__ == "__main__":
    main()
