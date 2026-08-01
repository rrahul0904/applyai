import io
import uuid

import pytest
from docx import Document
from sqlalchemy import create_engine, func, select
from sqlalchemy.exc import IntegrityError
from sqlalchemy.orm import Session

from app.core.config import Settings, get_settings
from app.core.outbox import publish_outbox_once
from app.core.queue import Task, TaskQueue
from app.core.storage import StorageObjectMetadata, get_object_storage
from app.durability_models import ResumeProcessingAttempt, ResumeUploadIntent, TaskOutbox
from app.main import app
from app.models import CandidateExperience, Resume, ResumeExtraction, ResumeVersion, User
from app.resumes.processor import PARSER_VERSION, process_resume_version
from tests.conftest import MemoryStorage


class DirectMemoryStorage(MemoryStorage):
    @property
    def supports_direct_upload(self) -> bool:
        return True

    def create_presigned_put(
        self,
        *,
        key: str,
        content_type: str,
        expires_in_seconds: int,
    ) -> str:
        del content_type, expires_in_seconds
        return f"https://s3.example.test/{key}"


class FailingQueue(TaskQueue):
    def enqueue(self, task: Task) -> None:
        del task
        raise RuntimeError("queue unavailable")


def docx_bytes() -> bytes:
    buffer = io.BytesIO()
    document = Document()
    document.add_paragraph("Candidate Name")
    document.add_paragraph("Senior Data Engineer")
    document.add_paragraph("Experience")
    document.add_paragraph("Senior Data Engineer | Example Labs | 2021 - Present")
    document.add_paragraph(
        "Built reliable data platforms and production pipelines for analytics, "
        "machine learning, governance, observability, reliability, privacy, security, "
        "cost controls, orchestration, testing, deployment, operations, and incident response."
    )
    document.add_paragraph("Skills")
    document.add_paragraph("Python, SQL, AWS, PostgreSQL, Airflow, Docker")
    document.save(buffer)
    return buffer.getvalue()


def test_replacement_uploads_version_one_master_resume(client, database_url):
    content = docx_bytes()
    first = client.post(
        "/api/v1/resumes",
        files={
            "file": (
                "resume.docx",
                io.BytesIO(content),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    second = client.post(
        "/api/v1/resumes",
        files={
            "file": (
                "resume-new.docx",
                io.BytesIO(content),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert first.status_code == 201
    assert second.status_code == 201
    assert first.json()["resume_id"] == second.json()["resume_id"]

    engine = create_engine(database_url)
    with Session(engine) as session:
        user = session.scalar(select(User).where(User.clerk_user_id == "clerk_user_a"))
        assert user is not None
        masters = list(
            session.scalars(
                select(Resume).where(Resume.user_id == user.id, Resume.is_master.is_(True))
            )
        )
        assert len(masters) == 1
        versions = list(
            session.scalars(
                select(ResumeVersion)
                .where(ResumeVersion.resume_id == masters[0].id)
                .order_by(ResumeVersion.version_number)
            )
        )
        assert [version.version_number for version in versions] == [1, 2]
    engine.dispose()


def test_database_prevents_second_master_resume(client, database_url):
    content = docx_bytes()
    response = client.post(
        "/api/v1/resumes",
        files={
            "file": (
                "resume.docx",
                io.BytesIO(content),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 201

    engine = create_engine(database_url)
    with Session(engine) as session:
        user = session.scalar(select(User).where(User.clerk_user_id == "clerk_user_a"))
        assert user is not None
        session.add(Resume(user_id=user.id, name="duplicate", is_master=True))
        with pytest.raises(IntegrityError):
            session.commit()
        session.rollback()
    engine.dispose()


def test_direct_upload_intent_is_owned_and_completion_creates_outbox(
    client,
    database_url,
    switch_user,
):
    storage = DirectMemoryStorage()
    durable_test_settings = Settings(
        task_queue_provider="sqs",
        sqs_queue_url="https://sqs.us-east-1.amazonaws.com/123456789012/test-resume",
    )
    app.dependency_overrides[get_object_storage] = lambda: storage
    app.dependency_overrides[get_settings] = lambda: durable_test_settings
    content = docx_bytes()

    intent_response = client.post(
        "/api/v1/resumes/upload-intents",
        json={
            "filename": "resume.docx",
            "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "file_size": len(content),
        },
    )
    assert intent_response.status_code == 200
    payload = intent_response.json()
    assert payload["upload_mode"] == "DIRECT_S3"
    assert payload["upload_url"].startswith("https://s3.example.test/")
    assert payload["upload_headers"]["x-amz-server-side-encryption"] == "AES256"

    version_id = uuid.UUID(payload["resume_version_id"])
    engine = create_engine(database_url)
    with Session(engine) as session:
        # A presign request is not a completed resume upload. Canonical versioning
        # begins only after S3 object verification succeeds.
        assert session.get(ResumeVersion, version_id) is None
        intent = session.scalar(
            select(ResumeUploadIntent).where(
                ResumeUploadIntent.resume_version_id == version_id
            )
        )
        assert intent is not None
        assert intent.status == "PENDING"
        storage.objects[intent.storage_key] = content
        storage.content_types[intent.storage_key] = intent.content_type
    engine.dispose()

    switch_user("clerk_user_b", "b@example.com")
    assert client.post(f"/api/v1/resumes/versions/{version_id}/upload-complete").status_code == 404
    switch_user("clerk_user_a", "a@example.com")

    completed = client.post(f"/api/v1/resumes/versions/{version_id}/upload-complete")
    assert completed.status_code == 200
    assert completed.json()["upload_status"] == "UPLOADED"
    assert completed.json()["processing_status"] == "QUEUED"

    # Completion is idempotent: retrying the callback returns the same version
    # without another version or outbox event.
    repeated = client.post(f"/api/v1/resumes/versions/{version_id}/upload-complete")
    assert repeated.status_code == 200
    assert repeated.json()["id"] == completed.json()["id"]

    engine = create_engine(database_url)
    with Session(engine) as session:
        version = session.get(ResumeVersion, version_id)
        assert version is not None
        assert version.upload_status == "UPLOADED"
        intent = session.scalar(
            select(ResumeUploadIntent).where(
                ResumeUploadIntent.resume_version_id == version_id
            )
        )
        assert intent is not None and intent.status == "COMPLETED"
        event = session.scalar(
            select(TaskOutbox).where(TaskOutbox.aggregate_id == version_id)
        )
        assert event is not None
        assert event.status == "PENDING"
        assert event.idempotency_key == f"resume-parse:{version_id}"
        assert session.scalar(
            select(func.count(ResumeVersion.id)).where(ResumeVersion.id == version_id)
        ) == 1
        assert session.scalar(
            select(func.count(TaskOutbox.id)).where(TaskOutbox.aggregate_id == version_id)
        ) == 1
    engine.dispose()


def test_upload_completion_rejects_object_size_mismatch(client, database_url):
    storage = DirectMemoryStorage()
    durable_test_settings = Settings(
        task_queue_provider="sqs",
        sqs_queue_url="https://sqs.us-east-1.amazonaws.com/123456789012/test-resume",
    )
    app.dependency_overrides[get_object_storage] = lambda: storage
    app.dependency_overrides[get_settings] = lambda: durable_test_settings

    intent_response = client.post(
        "/api/v1/resumes/upload-intents",
        json={
            "filename": "resume.docx",
            "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "file_size": 500,
        },
    )
    version_id = uuid.UUID(intent_response.json()["resume_version_id"])
    engine = create_engine(database_url)
    with Session(engine) as session:
        assert session.get(ResumeVersion, version_id) is None
        intent = session.scalar(
            select(ResumeUploadIntent).where(
                ResumeUploadIntent.resume_version_id == version_id
            )
        )
        assert intent is not None
        storage.objects[intent.storage_key] = b"short"
        storage.content_types[intent.storage_key] = intent.content_type
    engine.dispose()

    completed = client.post(f"/api/v1/resumes/versions/{version_id}/upload-complete")
    assert completed.status_code == 409
    assert completed.json()["error"]["code"] == "UPLOAD_SIZE_MISMATCH"

    engine = create_engine(database_url)
    with Session(engine) as session:
        assert session.get(ResumeVersion, version_id) is None
        intent = session.scalar(
            select(ResumeUploadIntent).where(
                ResumeUploadIntent.resume_version_id == version_id
            )
        )
        assert intent is not None and intent.status == "PENDING"
    engine.dispose()


def test_queue_failure_preserves_resume_and_retries_outbox(client, database_url):
    content = docx_bytes()
    uploaded = client.post(
        "/api/v1/resumes",
        files={
            "file": (
                "resume.docx",
                io.BytesIO(content),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert uploaded.status_code == 201
    version_id = uuid.UUID(uploaded.json()["id"])

    # Reset the already-published local-dev event so this test can exercise a queue
    # outage deterministically without coupling upload success to queue availability.
    engine = create_engine(database_url)
    with Session(engine) as session:
        event = session.scalar(select(TaskOutbox).where(TaskOutbox.aggregate_id == version_id))
        assert event is not None
        event.status = "PENDING"
        event.published_at = None
        event.available_at = func.now()
        session.commit()
    engine.dispose()

    claimed = publish_outbox_once(Settings(), queue=FailingQueue(), lock_owner="test-publisher")
    assert claimed == 1

    engine = create_engine(database_url)
    with Session(engine) as session:
        version = session.get(ResumeVersion, version_id)
        assert version is not None
        assert version.upload_status == "UPLOADED"
        event = session.scalar(select(TaskOutbox).where(TaskOutbox.aggregate_id == version_id))
        assert event is not None
        assert event.status == "PENDING"
        assert event.attempt_count == 1
        assert event.last_error == "RuntimeError"
    engine.dispose()


def test_resume_processing_redelivery_is_idempotent(client, database_url):
    content = docx_bytes()
    uploaded = client.post(
        "/api/v1/resumes",
        files={
            "file": (
                "resume.docx",
                io.BytesIO(content),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert uploaded.status_code == 201
    version_id = uuid.UUID(uploaded.json()["id"])

    # The development BackgroundTask may already have completed once. Explicit
    # redelivery must never create another extraction or successful attempt.
    process_resume_version(version_id, client.storage)
    process_resume_version(version_id, client.storage)

    engine = create_engine(database_url)
    with Session(engine) as session:
        extractions = list(
            session.scalars(
                select(ResumeExtraction).where(
                    ResumeExtraction.resume_version_id == version_id,
                    ResumeExtraction.parser_version == PARSER_VERSION,
                )
            )
        )
        attempts = list(
            session.scalars(
                select(ResumeProcessingAttempt).where(
                    ResumeProcessingAttempt.resume_version_id == version_id,
                    ResumeProcessingAttempt.parser_version == PARSER_VERSION,
                )
            )
        )
        assert len(extractions) == 1
        assert extractions[0].status in {"NEEDS_REVIEW", "COMPLETED"}
        assert len([attempt for attempt in attempts if attempt.status == "COMPLETED"]) == 1
    engine.dispose()


def test_failed_processing_can_retry_same_extraction(client, database_url):
    engine = create_engine(database_url)
    storage = MemoryStorage()
    with Session(engine) as session:
        user = User(clerk_user_id="retry-user", email="retry@example.test")
        session.add(user)
        session.flush()
        resume = Resume(user_id=user.id, name="retry", is_master=True)
        session.add(resume)
        session.flush()
        version = ResumeVersion(
            resume_id=resume.id,
            user_id=user.id,
            version_number=1,
            filename="resume.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            storage_key=f"candidate/{user.id}/resume/{resume.id}/retry.docx",
            file_size=7,
            upload_status="UPLOADED",
            processing_status="QUEUED",
        )
        session.add(version)
        session.commit()
        version_id = version.id
        storage.objects[version.storage_key] = b"invalid"
        storage.content_types[version.storage_key] = version.content_type
    engine.dispose()

    process_resume_version(version_id, storage)
    storage.objects[next(iter(storage.objects))] = docx_bytes()
    process_resume_version(version_id, storage)

    engine = create_engine(database_url)
    with Session(engine) as session:
        extraction_count = session.scalar(
            select(func.count(ResumeExtraction.id)).where(
                ResumeExtraction.resume_version_id == version_id
            )
        )
        attempts = list(
            session.scalars(
                select(ResumeProcessingAttempt)
                .where(ResumeProcessingAttempt.resume_version_id == version_id)
                .order_by(ResumeProcessingAttempt.attempt_number)
            )
        )
        version = session.get(ResumeVersion, version_id)
        assert extraction_count == 1
        assert [attempt.status for attempt in attempts] == ["FAILED", "COMPLETED"]
        assert version is not None and version.processing_status == "NEEDS_REVIEW"
    engine.dispose()


def test_resume_confirmation_completes_extraction_and_verifies_profile(client, database_url):
    content = docx_bytes()
    uploaded = client.post(
        "/api/v1/resumes",
        files={
            "file": (
                "resume.docx",
                io.BytesIO(content),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert uploaded.status_code == 201
    resume_id = uploaded.json()["resume_id"]

    response = client.post(
        f"/api/v1/resumes/{resume_id}/confirm",
        json={
            "headline": "Verified data engineer",
            "current_title": "Senior Data Engineer",
            "summary": "Candidate-reviewed profile.",
            "years_experience": 8,
            "target_roles": ["Staff Data Engineer"],
            "location_text": "Boston, MA",
            "work_modes": ["REMOTE"],
            "minimum_compensation": 180000,
            "experiences": [
                {
                    "company_name": "Example Labs",
                    "title": "Senior Data Engineer",
                    "description": "Reviewed by candidate",
                    "provenance": "DOCUMENT_EXTRACTED",
                }
            ],
            "education": [],
            "skills": [{"name": "Python", "provenance": "DOCUMENT_EXTRACTED"}],
        },
    )
    assert response.status_code == 200

    engine = create_engine(database_url)
    with Session(engine) as session:
        resume = session.get(Resume, uuid.UUID(resume_id))
        assert resume is not None
        version = session.scalar(
            select(ResumeVersion)
            .where(ResumeVersion.resume_id == resume.id)
            .order_by(ResumeVersion.version_number.desc())
        )
        assert version is not None and version.processing_status == "COMPLETED"
        extraction = session.scalar(
            select(ResumeExtraction).where(ResumeExtraction.resume_version_id == version.id)
        )
        assert extraction is not None and extraction.status == "COMPLETED"
        experience = session.scalar(select(CandidateExperience))
        assert experience is not None and experience.provenance == "USER_VERIFIED"
    engine.dispose()
